"""
Universal DOCX builder for insurance documents.

Renders any document from a declarative template definition (dict/JSON) + runtime data dict.
No document-specific code required — add new document types by defining a new template JSON.

Template schema:
  {
    "id": "unique-template-id",
    "title": "Document Title",
    "logo_key": "logo_path",          # optional: data key holding logo file path
    "company": { ... },               # optional overrides
    "sections": [                     # ordered list of sections
      {
        "type": "...",                # see SECTION TYPES below
        ... section-specific fields
      }
    ]
  }

SECTION TYPES:
  header        — company logo + name + address
  title         — large centred document title
  subtitle      — smaller centred subtitle
  notice        — boxed notice / warning paragraph
  paragraph     — free text paragraph
  heading       — section heading (level 1 or 2)
  divider       — horizontal rule
  key_value     — numbered label:value table (proposal fields)
  plan_select   — 2-col plan selection grid (e.g. Schengen/Non-Schengen)
  table         — data table with headers + rows from data key OR static rows
  signature     — signature block
  page_break    — page break
  declaration   — numbered declaration list

DATA BINDING:
  Any string value in the template can reference runtime data via {{ key }} syntax.
  Tables with rows_key pull their rows from data[rows_key] (list of dicts).
"""
from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ─── Defaults ────────────────────────────────────────────────────────────────
DEFAULT_FONT       = "Times New Roman"
DEFAULT_FONT_SIZE  = 11
COLOR_NAV          = "1F3864"   # Pragati dark navy
COLOR_NAV_LIGHT    = "E9EFF7"   # alternating row
COLOR_WHITE_STR    = "FFFFFF"
COLOR_NAV_RGB      = RGBColor(0x1F, 0x38, 0x64)
COLOR_WHITE_RGB    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK_RGB    = RGBColor(0x00, 0x00, 0x00)
COLOR_NOTICE_BG    = "EFF3FB"
COLOR_WARN_BG      = "FFF8E7"


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _interpolate(text: str, data: dict[str, Any]) -> str:
    """Replace {{ key }} placeholders with data values."""
    def replacer(m: re.Match) -> str:
        key = m.group(1).strip()
        val = data.get(key, "")
        return str(val) if val is not None else ""
    return re.sub(r"\{\{\s*(\w+)\s*\}\}", replacer, str(text))


def _hex(color: str) -> RGBColor:
    c = color.lstrip("#")
    return RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))


def _cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _remove_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBorders.append(el)
            tcPr.append(tcBorders)


def _para_fmt(p, space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = align


def _run(p, text: str, bold=False, italic=False,
         font=DEFAULT_FONT, size=DEFAULT_FONT_SIZE,
         color: RGBColor | None = None) -> None:
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


# ─── DocxBuilder ─────────────────────────────────────────────────────────────

class DocxBuilder:
    """
    Universal document builder.
    Usage:
        builder = DocxBuilder(template_def, data)
        buf = builder.build()   # returns io.BytesIO
    """

    def __init__(self, template: dict[str, Any], data: dict[str, Any]) -> None:
        self.tpl = template
        self.data = data
        self.font = template.get("font", DEFAULT_FONT)
        self.font_size = int(template.get("font_size", DEFAULT_FONT_SIZE))
        nav = template.get("nav_color", COLOR_NAV)
        self.color_nav = nav
        self.color_nav_rgb = _hex(nav)
        self.color_nav_light = template.get("nav_light_color", COLOR_NAV_LIGHT)

    def _iv(self, text: str) -> str:
        """Interpolate template text with runtime data."""
        return _interpolate(text, self.data)

    def _add_para(self, doc, text="", bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.LEFT, size=None,
                  color: RGBColor | None = None,
                  space_before=0, space_after=4) -> Any:
        p = doc.add_paragraph()
        _para_fmt(p, space_before=space_before, space_after=space_after, align=align)
        if text:
            _run(p, self._iv(text), bold=bold, italic=italic,
                 font=self.font, size=size or self.font_size, color=color)
        return p

    def _add_divider(self, doc) -> None:
        p = doc.add_paragraph()
        _para_fmt(p, space_before=2, space_after=2)
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), self.color_nav)
        border.append(bottom)
        p._p.get_or_add_pPr().append(border)

    def _add_notice(self, doc, text: str, bold=False, bg=COLOR_NOTICE_BG) -> None:
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        c = t.cell(0, 0)
        _cell_bg(c, bg)
        p = c.paragraphs[0]
        _para_fmt(p, space_before=3, space_after=3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _run(p, self._iv(text), bold=bold, font=self.font, size=self.font_size - 1)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _add_heading(self, doc, text: str, level=1) -> None:
        p = doc.add_paragraph()
        size = 14 if level == 1 else self.font_size
        _para_fmt(p, space_before=8 if level == 1 else 5, space_after=4)
        _run(p, self._iv(text), bold=True, font=self.font,
             size=size, color=self.color_nav_rgb)
        if level == 1:
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), self.color_nav)
            border.append(bottom)
            p._p.get_or_add_pPr().append(border)

    def _add_kv_table(self, doc, rows: list[dict], label_width=3.5) -> None:
        """Key-value table: number | label | value"""
        if not rows:
            return
        t = doc.add_table(rows=len(rows), cols=3)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, row_def in enumerate(rows):
            bg = self.color_nav_light if i % 2 == 0 else "FFFFFF"
            num   = self._iv(str(row_def.get("number", "")))
            label = self._iv(str(row_def.get("label", "")))
            value = self._iv(str(row_def.get("value", "")))
            c0, c1, c2 = t.rows[i].cells
            for c in (c0, c1, c2):
                _cell_bg(c, bg)
                c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                c.paragraphs[0].paragraph_format.space_before = Pt(2)
                c.paragraphs[0].paragraph_format.space_after = Pt(2)
            c0.width = Inches(0.35)
            c1.width = Inches(label_width)
            c2.width = Inches(6.46 - 0.35 - label_width)
            _run(c0.paragraphs[0], num, bold=True, font=self.font, size=self.font_size)
            _run(c1.paragraphs[0], label, font=self.font, size=self.font_size)
            _run(c2.paragraphs[0], value or "________________________________",
                 font=self.font, size=self.font_size)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _add_data_table(self, doc, headers: list[str], data_rows: list[list[str]],
                        col_widths: list[float] | None = None) -> None:
        if not headers:
            return
        n_cols = len(headers)
        n_rows = len(data_rows)
        t = doc.add_table(rows=1 + max(n_rows, 1), cols=n_cols)
        t.style = "Table Grid"

        # Auto column widths
        total = 6.46
        if col_widths and len(col_widths) == n_cols:
            widths = col_widths
        else:
            widths = [total / n_cols] * n_cols

        hdr_row = t.rows[0]
        for j, hdr in enumerate(headers):
            c = hdr_row.cells[j]
            _cell_bg(c, self.color_nav)
            c.width = Inches(widths[j])
            p = c.paragraphs[0]
            _para_fmt(p, space_before=3, space_after=3)
            _run(p, self._iv(str(hdr)), bold=True, font=self.font,
                 size=self.font_size - 1, color=COLOR_WHITE_RGB)

        for i in range(max(n_rows, 1)):
            bg = self.color_nav_light if i % 2 == 0 else "FFFFFF"
            row_data = data_rows[i] if i < n_rows else [""] * n_cols
            for j in range(n_cols):
                c = t.rows[i + 1].cells[j]
                _cell_bg(c, bg)
                c.width = Inches(widths[j])
                p = c.paragraphs[0]
                _para_fmt(p, space_before=2, space_after=2)
                val = self._iv(str(row_data[j]) if j < len(row_data) else "")
                _run(p, val, font=self.font, size=self.font_size)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _add_signature_block(self, doc, signatories: list[dict]) -> None:
        n = len(signatories)
        if not n:
            return
        t = doc.add_table(rows=2, cols=n)
        t.style = "Table Grid"
        for j, sig in enumerate(signatories):
            label = self._iv(str(sig.get("label", "")))
            name  = self._iv(str(sig.get("name", "")))
            tc = t.rows[0].cells[j]
            _cell_bg(tc, "F0F4FB")
            p = tc.paragraphs[0]
            _para_fmt(p, space_before=32, space_after=2)
            _run(p, name, font=self.font, size=self.font_size)
            bc = t.rows[1].cells[j]
            _cell_bg(bc, self.color_nav)
            pb = bc.paragraphs[0]
            _para_fmt(pb, space_before=2, space_after=2)
            _run(pb, label, bold=True, font=self.font,
                 size=self.font_size - 2, color=COLOR_WHITE_RGB)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def _add_header(self, doc, section_def: dict) -> None:
        """Company logo + name in a borderless 2-col table."""
        import os
        logo_path = self.data.get(
            section_def.get("logo_key", "logo_path"),
            section_def.get("logo_path", ""),
        )
        company_name = self._iv(section_def.get("company_name",
            self.tpl.get("company", {}).get("name", "PRAGATI INSURANCE PLC")))
        address_line = self._iv(section_def.get("address",
            self.tpl.get("company", {}).get("address",
            "20-21 Kawran Bazar, Dhaka-1215  |  Tel: 9133680-2  |  Fax: 880-2-55013694")))
        web_line = self._iv(section_def.get("web",
            self.tpl.get("company", {}).get("web",
            "info@pragatiinsurance.com  |  www.pragatiinsurance.com")))

        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        _remove_table_borders(t)

        logo_cell = t.rows[0].cells[0]
        logo_cell.width = Inches(1.4)
        text_cell = t.rows[0].cells[1]
        text_cell.width = Inches(5.06)

        lp = logo_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        if logo_path and os.path.isfile(str(logo_path)):
            lp.add_run().add_picture(str(logo_path), width=Inches(1.2))

        p_co = text_cell.paragraphs[0]
        p_co.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_fmt(p_co, space_before=0, space_after=2)
        _run(p_co, company_name, bold=True, font=self.font, size=19,
             color=self.color_nav_rgb)

        p_addr = text_cell.add_paragraph()
        _para_fmt(p_addr, space_before=0, space_after=0)
        _run(p_addr, address_line, font=self.font, size=9)

        p_web = text_cell.add_paragraph()
        _para_fmt(p_web, space_before=0, space_after=0)
        _run(p_web, web_line, font=self.font, size=9)

    def _add_plan_select(self, doc, section_def: dict) -> None:
        """2-col plan selection grid."""
        label = self._iv(str(section_def.get("label", "Plan Type:")))
        self._add_para(doc, label, space_before=4, space_after=2)

        options = section_def.get("options", [])
        if not options:
            return
        col_a = [o for o in options if o.get("col") == "A" or "col" not in o]
        col_b = [o for o in options if o.get("col") == "B"]
        # Build rows: header + one per option pair
        n_rows = 1 + max(len(col_a), len(col_b), 1)
        t = doc.add_table(rows=n_rows, cols=2)
        t.style = "Table Grid"
        # Header
        for j, hdr in enumerate(section_def.get("columns", ["Column A", "Column B"])):
            c = t.rows[0].cells[j]
            _cell_bg(c, self.color_nav)
            p = c.paragraphs[0]
            _para_fmt(p, space_before=2, space_after=2)
            _run(p, hdr, bold=True, font=self.font,
                 size=self.font_size, color=COLOR_WHITE_RGB)
        # Option rows — merge col_a and col_b by index
        all_left  = [self._iv(str(o.get("text", ""))) for o in col_a]
        all_right = [self._iv(str(o.get("text", ""))) for o in col_b]
        n = max(len(all_left), len(all_right))
        for i in range(n):
            bg = self.color_nav_light if i % 2 == 0 else "FFFFFF"
            left_txt  = all_left[i]  if i < len(all_left)  else ""
            right_txt = all_right[i] if i < len(all_right) else ""
            for j, txt in enumerate([left_txt, right_txt]):
                c = t.rows[i + 1].cells[j]
                _cell_bg(c, bg)
                p = c.paragraphs[0]
                _para_fmt(p, space_before=2, space_after=2)
                _run(p, txt, font=self.font, size=self.font_size)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _add_declaration(self, doc, items: list[str]) -> None:
        for i, item in enumerate(items, 1):
            p = self._add_para(doc, f"{i}.  {self._iv(item)}",
                               space_before=2, space_after=2)

    # ── Main build ────────────────────────────────────────────────────────────

    def build(self) -> io.BytesIO:
        doc = Document()

        # Page setup: A4
        for sec in doc.sections:
            sec.page_width   = Cm(21.0)
            sec.page_height  = Cm(29.7)
            sec.left_margin  = Cm(2.54)
            sec.right_margin = Cm(2.54)
            sec.top_margin   = Cm(1.6)
            sec.bottom_margin = Cm(1.9)

        doc.styles["Normal"].font.name = self.font
        doc.styles["Normal"].font.size = Pt(self.font_size)

        for section in self.tpl.get("sections", []):
            stype = section.get("type", "")

            if stype == "header":
                self._add_header(doc, section)

            elif stype == "divider":
                self._add_divider(doc)

            elif stype == "title":
                self._add_para(doc,
                    section.get("text", ""),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    size=section.get("size", 14),
                    color=self.color_nav_rgb,
                    space_before=section.get("space_before", 6),
                    space_after=section.get("space_after", 2))

            elif stype == "subtitle":
                self._add_para(doc,
                    section.get("text", ""),
                    bold=section.get("bold", False),
                    italic=section.get("italic", False),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    size=section.get("size", self.font_size),
                    space_before=section.get("space_before", 0),
                    space_after=section.get("space_after", 4))

            elif stype == "notice":
                bg = section.get("bg", COLOR_NOTICE_BG)
                self._add_notice(doc, section.get("text", ""),
                                 bold=section.get("bold", False), bg=bg)

            elif stype == "paragraph":
                align_map = {
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                }
                align = align_map.get(section.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                self._add_para(doc,
                    section.get("text", ""),
                    bold=section.get("bold", False),
                    italic=section.get("italic", False),
                    align=align,
                    size=section.get("size", self.font_size),
                    space_before=section.get("space_before", 0),
                    space_after=section.get("space_after", 4))

            elif stype == "heading":
                self._add_heading(doc, section.get("text", ""),
                                  level=section.get("level", 1))

            elif stype == "key_value":
                self._add_kv_table(doc,
                    section.get("rows", []),
                    label_width=section.get("label_width", 3.5))

            elif stype == "plan_select":
                self._add_plan_select(doc, section)

            elif stype == "table":
                headers = section.get("headers", [])
                rows_key = section.get("rows_key")
                static_rows = section.get("rows")
                col_widths = section.get("col_widths")

                if rows_key:
                    raw = self.data.get(rows_key, [])
                    cols = section.get("columns", [])
                    data_rows = [
                        [self._iv(str(r.get(c, ""))) for c in cols]
                        for r in (raw if isinstance(raw, list) else [])
                    ]
                elif static_rows:
                    data_rows = [
                        [self._iv(str(cell)) for cell in row]
                        for row in static_rows
                    ]
                else:
                    # rows from data key matching headers (auto-match)
                    data_rows = []

                self._add_data_table(doc, headers, data_rows, col_widths=col_widths)

            elif stype == "signature":
                self._add_signature_block(doc, section.get("signatories", []))

            elif stype == "page_break":
                doc.add_page_break()

            elif stype == "declaration":
                self._add_declaration(doc, section.get("items", []))

            elif stype == "footer":
                p = self._add_para(doc,
                    section.get("text", "Proposal Ref: {{ proposal_id }}   |   Generated: {{ generated_at }}"),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    italic=True,
                    size=9,
                    space_before=4,
                    space_after=0)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
