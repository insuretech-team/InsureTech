"""
DOCX renderer using python-docx + Jinja2.

Template content is a JSON descriptor describing document sections.
This approach lets Go store the template as a plain JSON string in the DB
(same `template_content` field) and the sidecar interprets it into a rich Word doc.

Template JSON schema example:
{
  "sections": [
    {"type": "title",       "text": "{{ policy_title }}"},
    {"type": "subtitle",    "text": "Policy #{{ policy_number }}"},
    {"type": "heading",     "text": "Policyholder Details", "level": 1},
    {"type": "paragraph",   "text": "Name: {{ policy_holder_name }}"},
    {"type": "key_value",   "rows": [
        {"label": "Policy Number", "value": "{{ policy_number }}"},
        {"label": "Start Date",    "value": "{{ start_date }}"},
        {"label": "End Date",      "value": "{{ end_date }}"}
    ]},
    {"type": "table", "headers": ["Description", "Qty", "Unit Price", "Amount"],
     "rows_key": "items", "columns": ["description", "quantity", "unit_price", "amount"]},
    {"type": "totals_table", "rows": [
        {"label": "Subtotal", "value": "{{ subtotal }}"},
        {"label": "Tax",      "value": "{{ tax }}"},
        {"label": "Total",    "value": "{{ total }}", "bold": true}
    ]},
    {"type": "paragraph",   "text": "Terms & Conditions", "style": "Heading 2"},
    {"type": "paragraph",   "text": "{{ terms }}"},
    {"type": "page_break"},
    {"type": "signature_block", "signatories": [
        {"label": "Authorized Signatory", "name": "{{ authorized_by }}"},
        {"label": "Policyholder",         "name": "{{ policy_holder_name }}"}
    ]}
  ]
}
"""

from __future__ import annotations

import io
import json
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Emu, Inches, Pt, RGBColor
from jinja2 import Environment, StrictUndefined, Undefined


# ─── Jinja2 env with lenient undefined ───────────────────────────────────────

class _SilentUndefined(Undefined):
    """Return empty string for missing template variables."""
    def __str__(self) -> str:
        return ""
    __iter__ = __str__
    __repr__ = __str__


_jinja_env = Environment(undefined=_SilentUndefined, autoescape=False)


def _render_str(template_str: str, data: dict) -> str:
    """Render a Jinja2 string expression against data."""
    try:
        return _jinja_env.from_string(str(template_str)).render(**data)
    except Exception:
        return str(template_str)


# ─── Colour palette ──────────────────────────────────────────────────────────

BRAND_PRIMARY   = RGBColor(0x0D, 0x47, 0xA1)   # deep blue
BRAND_ACCENT    = RGBColor(0x15, 0x65, 0xC0)   # medium blue
TABLE_HEADER_BG = "0D47A1"                       # hex string for XML
TABLE_ALT_ROW   = "EEF2F7"                       # light blue-grey
TOTALS_LABEL_BG = "E3EAF3"
TOTALS_FINAL_BG = "0D47A1"


# ─── XML helpers ─────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set cell background color via OOXML shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, border_spec: dict | None = None) -> None:
    """Set cell borders. border_spec: {"top": "single", ...}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    sides = border_spec or {"top": "single", "bottom": "single", "left": "single", "right": "single"}
    for side, style in sides.items():
        el = OxmlElement(f"w:{'left' if side == 'start' else side}")
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_page_number(footer) -> None:
    """Insert 'Page X of Y' field into a paragraph."""
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Page ")
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)
    para.add_run(" of ")
    run2 = para.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    run2._r.append(fldChar3)
    instrText2 = OxmlElement("w:instrText")
    instrText2.text = "NUMPAGES"
    run2._r.append(instrText2)
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run2._r.append(fldChar4)


# ─── Section renderers ───────────────────────────────────────────────────────

def _render_title(doc: Document, section: dict, data: dict) -> None:
    text = _render_str(section.get("text", ""), data)
    p = doc.add_paragraph(style="Title")
    run = p.add_run(text)
    run.font.color.rgb = BRAND_PRIMARY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _render_subtitle(doc: Document, section: dict, data: dict) -> None:
    text = _render_str(section.get("text", ""), data)
    p = doc.add_paragraph(style="Subtitle")
    run = p.add_run(text)
    run.font.color.rgb = BRAND_ACCENT
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _render_heading(doc: Document, section: dict, data: dict) -> None:
    text = _render_str(section.get("text", ""), data)
    level = int(section.get("level", 1))
    level = max(1, min(level, 9))
    p = doc.add_heading(text, level=level)
    # Color brand blue
    for run in p.runs:
        run.font.color.rgb = BRAND_PRIMARY


def _render_paragraph(doc: Document, section: dict, data: dict) -> None:
    text = _render_str(section.get("text", ""), data)
    style = section.get("style")
    if style:
        try:
            p = doc.add_paragraph(style=style)
        except Exception:
            p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    if section.get("bold"):
        run.bold = True
    if section.get("italic"):
        run.italic = True
    if section.get("font_size"):
        run.font.size = Pt(float(section["font_size"]))
    align = section.get("align", "").lower()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _render_key_value(doc: Document, section: dict, data: dict) -> None:
    """Two-column key-value table."""
    rows_spec = section.get("rows", [])
    if not rows_spec:
        return
    table = doc.add_table(rows=len(rows_spec), cols=2)
    table.style = "Table Grid"
    col_widths = [Inches(2.2), Inches(4.0)]
    for i, row_spec in enumerate(rows_spec):
        row = table.rows[i]
        label = _render_str(row_spec.get("label", ""), data)
        value = _render_str(row_spec.get("value", ""), data)
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = col_widths[0]
        vc.width = col_widths[1]
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.color.rgb = BRAND_PRIMARY
        vc.paragraphs[0].add_run(value)
        _set_cell_bg(lc, TOTALS_LABEL_BG)
        _set_cell_border(lc)
        _set_cell_border(vc)
    doc.add_paragraph()


def _render_table(doc: Document, section: dict, data: dict) -> None:
    """Styled data table with branded header row."""
    headers = section.get("headers", [])
    rows_key = section.get("rows_key", "items")
    columns = section.get("columns", [])
    items = data.get(rows_key, [])
    if not isinstance(items, list):
        items = []

    if not headers:
        return

    n_cols = len(headers)
    n_rows = len(items) + 1  # +1 for header
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for col_i, hdr_text in enumerate(headers):
        cell = hdr_row.cells[col_i]
        _set_cell_bg(cell, TABLE_HEADER_BG)
        _set_cell_border(cell)
        run = cell.paragraphs[0].add_run(str(hdr_text))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for row_i, item in enumerate(items):
        row = table.rows[row_i + 1]
        if isinstance(item, dict):
            row_data = item
        else:
            row_data = {}
        alt = (row_i % 2 == 1)
        for col_i, col_key in enumerate(columns):
            cell = row.cells[col_i]
            _set_cell_border(cell)
            if alt:
                _set_cell_bg(cell, TABLE_ALT_ROW)
            val = str(row_data.get(col_key, ""))
            cell.paragraphs[0].add_run(val)

    doc.add_paragraph()


def _render_totals_table(doc: Document, section: dict, data: dict) -> None:
    """Right-aligned totals block."""
    rows_spec = section.get("rows", [])
    if not rows_spec:
        return
    table = doc.add_table(rows=len(rows_spec), cols=2)
    table.style = "Table Grid"
    for i, row_spec in enumerate(rows_spec):
        row = table.rows[i]
        label = _render_str(row_spec.get("label", ""), data)
        value = _render_str(row_spec.get("value", ""), data)
        is_final = bool(row_spec.get("bold", False))
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = Inches(4.5)
        vc.width = Inches(1.8)
        lr = lc.paragraphs[0].add_run(label)
        vr = vc.paragraphs[0].add_run(value)
        if is_final:
            lr.bold = True
            vr.bold = True
            lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            vr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(lc, TOTALS_FINAL_BG)
            _set_cell_bg(vc, TOTALS_FINAL_BG)
        else:
            _set_cell_bg(lc, TOTALS_LABEL_BG)
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_cell_border(lc)
        _set_cell_border(vc)
    doc.add_paragraph()


def _render_page_break(doc: Document, section: dict, data: dict) -> None:
    doc.add_page_break()


def _render_signature_block(doc: Document, section: dict, data: dict) -> None:
    """Signature lines at the bottom."""
    signatories = section.get("signatories", [])
    if not signatories:
        return
    doc.add_paragraph()
    n = len(signatories)
    table = doc.add_table(rows=2, cols=n)
    for col_i, sig in enumerate(signatories):
        label = _render_str(sig.get("label", ""), data)
        name  = _render_str(sig.get("name", ""), data)
        # Signature line row
        cell_line = table.rows[0].cells[col_i]
        p = cell_line.paragraphs[0]
        p.add_run("_" * 30)
        # Label+name row
        cell_label = table.rows[1].cells[col_i]
        pl = cell_label.paragraphs[0]
        r = pl.add_run(f"{label}\n{name}" if name else label)
        r.font.size = Pt(9)
        r.font.color.rgb = BRAND_ACCENT


def _render_horizontal_rule(doc: Document, section: dict, data: dict) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0D47A1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─── Section dispatcher ───────────────────────────────────────────────────────

_SECTION_RENDERERS = {
    "title":           _render_title,
    "subtitle":        _render_subtitle,
    "heading":         _render_heading,
    "paragraph":       _render_paragraph,
    "key_value":       _render_key_value,
    "table":           _render_table,
    "totals_table":    _render_totals_table,
    "page_break":      _render_page_break,
    "signature_block": _render_signature_block,
    "hr":              _render_horizontal_rule,
}


# ─── Public API ───────────────────────────────────────────────────────────────

def render_docx(
    template_content: str,
    data: dict,
    title: str = "",
    author: str = "InsureTech",
    subject: str = "",
) -> io.BytesIO:
    """
    Render a DOCX from a JSON template descriptor + data dict.

    template_content may be:
      - A JSON object with a "sections" array (preferred, structured)
      - A plain Jinja2 string → treated as a single paragraph block

    Returns a BytesIO buffer containing the .docx bytes.
    """
    doc = Document()

    # ── Core page setup ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Emu(12_192_000)   # A4 width  (8.27 in)
    section.page_height = Emu(15_874_560)   # A4 height (11.69 in)
    margin = Inches(1.0)
    section.left_margin   = margin
    section.right_margin  = margin
    section.top_margin    = margin
    section.bottom_margin = margin

    # ── Document core properties ──────────────────────────────────────────────
    core_props = doc.core_properties
    core_props.title   = _render_str(title, data) if title else _render_str("{{ policy_title|default('Document') }}", data)
    core_props.author  = author
    core_props.subject = subject

    # ── Footer with page numbers ──────────────────────────────────────────────
    footer = section.footer
    _add_page_number(footer)

    # ── Parse template ────────────────────────────────────────────────────────
    template_spec: dict = {}
    try:
        template_spec = json.loads(template_content)
    except (json.JSONDecodeError, TypeError):
        # Treat as raw Jinja2 text — render each line as paragraph
        rendered_text = _render_str(template_content, data)
        for line in rendered_text.splitlines():
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    sections = template_spec.get("sections", [])
    if not sections:
        # Flat key-value dump
        for k, v in data.items():
            if k.startswith("_"):
                continue
            doc.add_paragraph(f"{k}: {v}")
    else:
        for sec in sections:
            sec_type = str(sec.get("type", "paragraph")).lower()
            renderer = _SECTION_RENDERERS.get(sec_type, _render_paragraph)
            renderer(doc, sec, data)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
