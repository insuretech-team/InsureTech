"""
Overseas Mediclaim Proposal Form – Pragati Insurance PLC
Matches the exact format, fonts and layout of the KBank reference DOCX files.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ─── Constants ────────────────────────────────────────────────────────────────
FONT_BODY      = "Times New Roman"
FONT_BANGLA    = "SutonnyMJ"
PT_COMPANY     = 19
PT_TITLE       = 14
PT_BODY        = 11
PT_SMALL       = 9
PT_NOTICE      = 10
COLOR_HEADER_BG = "1F3864"   # dark navy – Pragati brand
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK     = RGBColor(0x00, 0x00, 0x00)
COLOR_DARK      = RGBColor(0x1F, 0x38, 0x64)
COLOR_LIGHT_BG  = "E9EFF7"   # light blue-grey for alternating rows


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, **borders) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        style = borders.get(edge, {"val": "single", "sz": "4", "color": "B0B8C9"})
        el = OxmlElement(f"w:{edge}")
        for k, v in style.items():
            el.set(qn(f"w:{k}"), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _run(para, text: str, bold=False, italic=False, font=FONT_BODY,
         size=PT_BODY, color: RGBColor | None = None) -> None:
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _para(doc, text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          font=FONT_BODY, size=PT_BODY, space_before=0, space_after=4,
          color: RGBColor | None = None) -> Any:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        _run(p, text, bold=bold, italic=italic, font=font, size=size, color=color)
    return p


def _heading(doc, text: str, level=1) -> Any:
    p = _para(doc, space_before=6, space_after=3)
    _run(p, text, bold=True, font=FONT_BODY,
         size=PT_TITLE if level == 1 else PT_BODY,
         color=COLOR_DARK)
    return p


def _divider(doc) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3864")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def _notice_box(doc, text: str, bold=False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "EFF3FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, bold=bold, font=FONT_BODY, size=PT_NOTICE)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _kv_table(doc, rows: list[tuple[str, str]]) -> None:
    """3-column table: # | label | value"""
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths: ~0.4" | 3.0" | 3.0"
    for i, row in enumerate(rows):
        num_text, label_text, value_text = row
        bg = COLOR_LIGHT_BG if i % 2 == 0 else "FFFFFF"

        c0, c1, c2 = table.rows[i].cells

        c0.width = Inches(0.35)
        c1.width = Inches(3.05)
        c2.width = Inches(3.05)

        for c in (c0, c1, c2):
            _set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        _run(p0, num_text, bold=True, font=FONT_BODY, size=PT_BODY)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        _run(p1, label_text, font=FONT_BODY, size=PT_BODY)

        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        _run(p2, value_text or "____________________________", font=FONT_BODY, size=PT_BODY)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _data_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    n_cols = len(headers)
    n_rows = len(rows)
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        _set_cell_bg(cell, COLOR_HEADER_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        _run(p, hdr, bold=True, font=FONT_BODY, size=PT_BODY, color=COLOR_WHITE)

    # Data rows
    for i, row in enumerate(rows):
        bg = COLOR_LIGHT_BG if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _run(p, str(val) if val else "____________________", font=FONT_BODY, size=PT_BODY)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _sig_table(doc, signatories: list[tuple[str, str]]) -> None:
    n = len(signatories)
    table = doc.add_table(rows=2, cols=n)
    table.style = "Table Grid"
    for j, (label, name) in enumerate(signatories):
        # Top cell: name / blank line
        tc = table.rows[0].cells[j]
        _set_cell_bg(tc, "F0F4FB")
        p = tc.paragraphs[0]
        p.paragraph_format.space_before = Pt(30)
        p.paragraph_format.space_after = Pt(2)
        _run(p, name or "", font=FONT_BODY, size=PT_BODY)
        # Bottom cell: label
        bc = table.rows[1].cells[j]
        _set_cell_bg(bc, COLOR_HEADER_BG)
        pb = bc.paragraphs[0]
        pb.paragraph_format.space_before = Pt(2)
        pb.paragraph_format.space_after = Pt(2)
        _run(pb, label, bold=True, font=FONT_BODY, size=PT_SMALL, color=COLOR_WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─── Main generator ──────────────────────────────────────────────────────────

def generate(data: dict[str, Any]) -> io.BytesIO:
    doc = Document()

    # ── Page Setup: A4, Pragati margins ──
    for section in doc.sections:
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
        section.top_margin    = Cm(1.6)
        section.bottom_margin = Cm(1.9)

    # ── Default paragraph font ──
    doc.styles["Normal"].font.name = FONT_BODY
    doc.styles["Normal"].font.size = Pt(PT_BODY)

    g = data.get  # shorthand

    # ════════════════════════════════════════════
    # HEADER – Logo + Company branding
    # ════════════════════════════════════════════
    LOGO_PATH = data.get(
        "logo_path",
        r"E:\Projects\InsureTech\web_shared\insurers\pragati_logo.png",
    )
    import os as _os

    # Logo + company name side by side in a borderless table
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = "Table Grid"
    # Remove all borders on the header table
    for row in hdr_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBorders.append(el)
            tcPr.append(tcBorders)

    logo_cell = hdr_table.rows[0].cells[0]
    logo_cell.width = Inches(1.4)
    text_cell = hdr_table.rows[0].cells[1]
    text_cell.width = Inches(5.06)

    # Insert logo
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after = Pt(0)
    if _os.path.isfile(LOGO_PATH):
        run_logo = logo_para.add_run()
        run_logo.add_picture(LOGO_PATH, width=Inches(1.2))

    # Company name + address in right cell
    p_co = text_cell.paragraphs[0]
    p_co.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_co.paragraph_format.space_before = Pt(0)
    p_co.paragraph_format.space_after = Pt(2)
    _run(p_co, "PRAGATI INSURANCE PLC", bold=True, font=FONT_BODY, size=PT_COMPANY, color=COLOR_DARK)

    p_addr = text_cell.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_addr.paragraph_format.space_before = Pt(0)
    p_addr.paragraph_format.space_after = Pt(0)
    _run(p_addr, "20-21 Kawran Bazar, Dhaka-1215  |  Tel: 9133680-2  |  Fax: 880-2-55013694",
         font=FONT_BODY, size=PT_SMALL)

    p_web = text_cell.add_paragraph()
    p_web.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_web.paragraph_format.space_before = Pt(0)
    p_web.paragraph_format.space_after = Pt(0)
    _run(p_web, "info@pragatiinsurance.com  |  www.pragatiinsurance.com",
         font=FONT_BODY, size=PT_SMALL)

    _divider(doc)

    # ── Document Title ──
    _para(doc, "PROPOSAL FORM FOR OVERSEAS MEDICLAIM POLICY",
          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          font=FONT_BODY, size=PT_TITLE, space_before=6, space_after=2, color=COLOR_DARK)

    _para(doc, "(BUSINESS AND HOLIDAYS)",
          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_BODY, size=PT_BODY,
          space_before=0, space_after=2)

    _para(doc, "(To be submitted in original with two copies)  |  (Available to persons aged 6 months to 79 years)",
          italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_BODY, size=PT_SMALL,
          space_before=0, space_after=6)

    _divider(doc)

    # ── Notice box ──
    _notice_box(doc,
        "THE OVERSEAS MEDICLAIM POLICY PROVIDES INDEMNITY FOR EXPENSES INCURRED FOR MEDICAL TREATMENT "
        "TO THE INSURED PERSON WHO TRAVELS ABROAD AS CORPORATE CLIENT, FOR ILLNESS, DISEASES CONTRACTED "
        "OR INJURY SUSTAINED DURING OVERSEAS TRAVEL AND WHICH IS PRIMARILY IN THE NATURE OF AN EMERGENCY "
        "AND WHICH IS NECESSARY TO BE UNDERTAKEN IMMEDIATELY, WITHOUT WHICH THE PROPOSER IS NOT ABLE TO "
        "LEAVE THE OVERSEAS COUNTRY UNDER MEDICAL ADVICE. THE ATTENTION OF THE PROPOSER IS DRAWN TO ITEM "
        "II (MEDICAL HISTORY) OF THE PROPOSAL FORM, ESPECIALLY IN RELATION TO PREVIOUS TREATMENT OF "
        "ILLNESS BY THE PROPOSER.",
        bold=True)

    _notice_box(doc,
        "THE PROPOSAL FORM SHOULD BE COMPLETED TO THE BEST OF YOUR KNOWLEDGE AND BELIEF AND ALL MATERIAL "
        "FACTS * SHOULD BE DISCLOSED. FAILURE TO DO SO MAY NULLIFY COVER UNDER ANY POLICY ISSUED.\n"
        "* A material fact is one that is likely to influence the Insurer's acceptance or assessment of the "
        "proposal. Consult the Corporation/Company if you are in any doubt as to what constitutes a material fact.")

    _divider(doc)

    # ════════════════════════════════════════════
    # SECTION I – PROPOSER DETAILS
    # ════════════════════════════════════════════
    _heading(doc, "I.  PROPOSER DETAILS", level=1)

    _kv_table(doc, [
        ("1.", "Name and status of the proposer (in block letters) as stated in the passport\n"
               "State whether Mr. / Mrs. / Miss / Master",
               g("proposer_name", "")),
        ("2.", "Residence Address",                         g("residence_address", "")),
        ("3.", "Residence Telephone No. & Mobile No.",      g("phone_number", "")),
        ("4.", "Proposer's Actual Occupation (specify)",    g("occupation", "")),
        ("5.", "Office Name and Address, if any",           g("office_name_address", "")),
        ("6.", "Office Telephone No.",                      g("office_telephone", "")),
        ("7.", "Age (in completed years)",                  g("age", "")),
        ("8.", "Passport Number (copy attached)",           g("passport_number", "")),
    ])

    # Plan type – special 2-column table
    _para(doc, "9.  Plan Type:", bold=False, font=FONT_BODY, size=PT_BODY,
          space_before=4, space_after=2)

    plan_table = doc.add_table(rows=3, cols=2)
    plan_table.style = "Table Grid"

    plans = [
        ("Schengen Countries", "Non-Schengen Countries"),
        ("Worldwide (excl. USA & Canada)  —  Plan A:  " + g("plan_schengen_a", "☐"),
         "Worldwide (excl. USA & Canada)  —  Plan A:  " + g("plan_non_schengen_a", "☐")),
        ("Worldwide (incl. USA & Canada)  —  Plan B:  " + g("plan_schengen_b", "☐"),
         "Worldwide (incl. USA & Canada)  —  Plan B:  " + g("plan_non_schengen_b", "☐")),
    ]
    for ri, (left, right) in enumerate(plans):
        c0 = plan_table.rows[ri].cells[0]
        c1 = plan_table.rows[ri].cells[1]
        if ri == 0:
            _set_cell_bg(c0, COLOR_HEADER_BG)
            _set_cell_bg(c1, COLOR_HEADER_BG)
            p0 = c0.paragraphs[0]; p1 = c1.paragraphs[0]
            _run(p0, left,  bold=True, font=FONT_BODY, size=PT_BODY, color=COLOR_WHITE)
            _run(p1, right, bold=True, font=FONT_BODY, size=PT_BODY, color=COLOR_WHITE)
        else:
            bg = COLOR_LIGHT_BG if ri % 2 == 1 else "FFFFFF"
            _set_cell_bg(c0, bg); _set_cell_bg(c1, bg)
            p0 = c0.paragraphs[0]; p1 = c1.paragraphs[0]
            _run(p0, left,  font=FONT_BODY, size=PT_BODY)
            _run(p1, right, font=FONT_BODY, size=PT_BODY)
        for c in (c0, c1):
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _kv_table(doc, [
        ("10.", "Purpose of Trip\n(Official / Holiday – conducted tour / Holiday – individual)",
                g("trip_purpose", "")),
        ("11.", "Proposed date of departure from Bangladesh\n(No extension can be granted)",
                g("departure_date", "")),
        ("12.", "Number of days stay outside Bangladesh\n(No extension can be granted)",
                g("days_abroad", "")),
        ("13.", "Itinerary (countries and places to be visited with approximate days at each place)",
                g("itinerary", "")),
        ("14.", "Name, Address & Registration No. of usual physician\n"
               "Telephone No. (Consulting Room / Office / Residence)",
                g("physician_name_address", "") + "\n" + g("physician_telephone", "")),
    ])

    # ════════════════════════════════════════════
    # SECTION II – MEDICAL HISTORY
    # ════════════════════════════════════════════
    doc.add_page_break()

    _heading(doc, "II.  MEDICAL HISTORY  —  TO BE COMPLETED BY THE PROPOSER / SPOUSE", level=1)

    _para(doc, "PLEASE ANSWER THE FOLLOWING QUESTIONS IN YES OR NO "
               "(A DASH IS NOT SUFFICIENT) AND GIVE FULL DETAILS.",
          bold=True, font=FONT_BODY, size=PT_BODY, space_before=0, space_after=6)

    _kv_table(doc, [
        ("1.", "Are you in good health and free from physical and mental disease or infirmity?",
               g("q1_good_health", "")),
        ("2(a).", "Any nervous, mental or psychiatric disease, slipped disc or other spinal disorder, "
                  "fainting episode, blackout, fit or paralysis of any kind?",
                  g("q2a_nervous", "")),
        ("2(b).", "High blood pressure, heart diseases including ischemic heart disease, piles, varicose "
                  "veins, other circulatory disorders or rheumatic fever?",
                  g("q2b_heart", "")),
        ("2(c).", "Hernia, any rheumatic or joint disease, urinary disease or diabetes?",
                  g("q2c_hernia", "")),
        ("2(d).", "Any respiratory or allergic disease, or any disorder of the stomach, bowel or gallbladder?",
                  g("q2d_respiratory", "")),
        ("2(e).", "Any other complaint requiring specialist's consultation or surgical or hospital "
                  "treatment or investigations?",
                  g("q2e_specialist", "")),
        ("2(f).", "Any complaint or tendency that may necessitate such consultation or treatment in the future?",
                  g("q2f_future", "")),
        ("3.", "Are there any additional facts affecting the proposed insurance which should be disclosed to Insurers?",
               g("q3_additional_facts", "")),
        ("4.", "Have you any intention of engaging in winter sports or pastimes rendering you liable to personal injury?",
               g("q4_winter_sports", "")),
    ])

    # Illness history table
    _para(doc, "5.  Give particulars of any other illness, disease or accident sustained during the "
               "12 months preceding the first day of Insurance:",
          font=FONT_BODY, size=PT_BODY, space_before=4, space_after=4)

    illness_rows = g("illness_history", [])
    if not illness_rows:
        illness_rows = [{}, {}, {}]
    while len(illness_rows) < 3:
        illness_rows.append({})

    _data_table(doc,
        ["Nature of Illness / Disease and Treatment Received",
         "Date First Treated",
         "Name & Address of Attending Practitioner / Surgeon with Telephone No."],
        [[r.get("nature_of_illness", ""), r.get("date_first_treated", ""),
          r.get("practitioner_details", "")] for r in illness_rows]
    )

    # Known ailments
    _para(doc, "Please give details of any ailment, sickness or injury which may require medical "
               "attention whilst on tour abroad:",
          font=FONT_BODY, size=PT_BODY, space_before=4, space_after=4)

    _kv_table(doc, [
        ("1.", "", g("known_ailment_1", "")),
        ("2.", "", g("known_ailment_2", "")),
        ("3.", "", g("known_ailment_3", "")),
        ("4.", "", g("known_ailment_4", "")),
    ])

    # ════════════════════════════════════════════
    # DECLARATION
    # ════════════════════════════════════════════
    doc.add_page_break()

    _heading(doc, "I HEREBY DECLARE THAT:", level=2)

    declarations = [
        "I will not be travelling against the advice of a physician.",
        "I am not on a waiting list for any medical treatment.",
        "I will not be travelling for the purpose of obtaining medical treatment.",
        "I have not received a terminal prognosis for a medical condition before this day.",
    ]
    for i, d in enumerate(declarations, 1):
        p = _para(doc, f"{i}.  {d}", font=FONT_BODY, size=PT_BODY,
                  space_before=2, space_after=2)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _notice_box(doc,
        "I further declare and warrant that the above statements are true and complete. "
        "I consent to the insurers seeking medical information from any doctor who has at any time attended "
        "concerning anything which affects my physical or mental health, and I authorise the giving of such "
        "information as Van Ameyde UK Ltd. / Specialty Assist Ltd. and/or their Program Medical Advisor may "
        "require. I agree that this proposal shall form the basis of the contract should the insurance be effected.\n\n"
        "I am willing to accept the Policy, subject to the terms, exceptions and conditions prescribed by "
        "Corporation/Company therein.")

    _sig_table(doc, [
        ("Signature", g("proposer_name", "")),
        ("Place",     g("signature_place", "")),
        ("Date (DD / MM / YY)", g("signature_date", "")),
    ])

    _divider(doc)

    # ── Schengen Countries ──
    _para(doc, "List of Schengen Countries", bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_BODY, size=PT_BODY,
          space_before=6, space_after=2, color=COLOR_DARK)

    _para(doc,
          "Austria, Belgium, Denmark, Finland, France, Germany, Iceland, Italy, Greece, Luxembourg, "
          "Netherlands, Norway, Portugal, Spain, Sweden, Estonia, Latvia, Lithuania, Poland, Czech Republic, "
          "Slovakia, Hungary, Slovenia, Malta, Cyprus, Switzerland and Liechtenstein",
          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_BODY, size=PT_SMALL,
          space_before=0, space_after=6)

    _divider(doc)

    # ── Product Benefits ──
    _para(doc, "OVERSEAS MEDICLAIM POLICY (TRAVEL INSURANCE) — PRODUCT BENEFITS & LIMITATIONS",
          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_BODY, size=PT_BODY,
          space_before=6, space_after=4, color=COLOR_DARK)

    benefits = g("product_benefits", [
        {"number": "01.", "benefit": "Medical Expenses & Hospitalization abroad (Worldwide excl. USA/Canada)", "limit": "US$ 50,000 – Excess USD 100"},
        {"number": "02.", "benefit": "Medical Expenses & Hospitalization abroad (Worldwide incl. USA/Canada)", "limit": "US$ 100,000 – Excess USD 100"},
        {"number": "03.", "benefit": "Medical Expenses & Hospitalization for Schengen Countries",             "limit": "Euro 30,000 – Nil deductible"},
        {"number": "04.", "benefit": "Transport or Repatriation in case of Illness or Accident",              "limit": "Actual Expenses"},
        {"number": "05.", "benefit": "Emergency Dental Care",                                                 "limit": "US$ 500 – Excess US$ 50"},
        {"number": "06.", "benefit": "Repatriation of Family Member Travelling with the Insured",             "limit": "Actual Expenses"},
        {"number": "07.", "benefit": "Repatriation of Mortal Remains",                                       "limit": "Actual Expenses"},
        {"number": "08.", "benefit": "Travel of one immediate family member",                                 "limit": "US$ 100/day – Max US$ 1,000"},
        {"number": "09.", "benefit": "Emergency return home following death of a close family member",        "limit": "Actual Expenses"},
    ])

    _data_table(doc,
        ["#", "Benefit / Coverage", "Coverage Limit"],
        [[b.get("number", ""), b.get("benefit", ""), b.get("limit", "")] for b in benefits])

    _notice_box(doc,
        "NOTE: THE COMPANY WILL NOT BE LIABLE TO PROVIDE ANY ASSISTANCE WHICH ARISES DIRECTLY OR "
        "INDIRECTLY FROM ANY PRE-EXISTING MEDICAL CONDITION, SUICIDE OR ATTEMPTED SUICIDE, MENTAL "
        "ILLNESS, PREGNANCY OR CHILDBIRTH.",
        bold=True)

    # ── Footer ref ──
    _divider(doc)
    _para(doc,
          f"Proposal Ref: {g('proposal_id', '')}   |   Generated: {g('generated_at', '')}",
          align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_BODY, size=PT_SMALL,
          italic=True, space_before=4, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
