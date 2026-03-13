#!/usr/bin/env python3
"""
Convert EXECUTIVE_SUMMARY.md to DOCX with professional formatting.
Simplified version of todocx.py for the executive summary.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def convert_exec_summary_to_docx(md_file, docx_file):
    """Convert Executive Summary markdown to DOCX."""
    
    print(f"Reading: {md_file}")
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process line by line
    lines = content.split('\n')
    in_table = False
    table_data = []
    
    for line in lines:
        stripped = line.strip()
        
        # Skip empty lines in tables
        if in_table and not stripped:
            # Process accumulated table
            if table_data:
                process_table(doc, table_data)
                table_data = []
            in_table = False
            continue
        
        # Table detection
        if '|' in stripped and stripped.startswith('|'):
            in_table = True
            # Skip separator rows
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                continue
            table_data.append(stripped)
            continue
        
        # If we were in a table and now hit non-table content
        if in_table and not stripped.startswith('|'):
            if table_data:
                process_table(doc, table_data)
                table_data = []
            in_table = False
        
        # Heading detection
        if stripped.startswith('#'):
            level = len(re.match(r'^#+', stripped).group())
            text = re.sub(r'^#+\s*', '', stripped)
            
            if level == 1:
                # Title (center, bold, 18pt)
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(18)
                    run.font.bold = True
            elif level == 2:
                # Main section (bold, 16pt)
                p = doc.add_heading(text, level=1)
                for run in p.runs:
                    run.font.size = Pt(16)
            elif level == 3:
                # Subsection (bold, 14pt)
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.size = Pt(14)
            else:
                # Lower level headings
                p = doc.add_heading(text, level=min(level, 3))
        
        # Bold/Italic text
        elif '**' in stripped or '*' in stripped:
            p = doc.add_paragraph()
            process_inline_formatting(p, stripped)
        
        # Bullet lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(text, style='List Number')
        
        # Horizontal rule
        elif stripped == '---':
            doc.add_page_break()
        
        # Regular paragraph
        elif stripped:
            p = doc.add_paragraph()
            process_inline_formatting(p, stripped)
    
    # Process any remaining table
    if table_data:
        process_table(doc, table_data)
    
    # Save document
    print(f"Saving: {docx_file}")
    doc.save(docx_file)
    print(f"✓ Conversion complete!")


def process_table(doc, table_data):
    """Process markdown table and add to document."""
    if not table_data:
        return
    
    # Parse table rows
    rows = []
    for row in table_data:
        cells = [cell.strip() for cell in row.strip('|').split('|')]
        rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_data
            
            # Style header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    doc.add_paragraph()  # Add space after table


def process_inline_formatting(paragraph, text):
    """Process inline formatting (bold, italic, code)."""
    # Split by formatting markers
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            # Regular text
            paragraph.add_run(part)


def main():
    script_dir = Path(__file__).parent
    md_file = script_dir / "EXECUTIVE_SUMMARY.md"
    docx_file = script_dir / "EXECUTIVE_SUMMARY.docx"
    
    print("="*60)
    print("Executive Summary Markdown to DOCX Converter")
    print("="*60)
    print()
    
    if not md_file.exists():
        print(f"❌ Error: {md_file} not found")
        return
    
    convert_exec_summary_to_docx(md_file, docx_file)
    
    print()
    print("="*60)
    print(f"✅ Output: {docx_file}")
    print("="*60)


if __name__ == "__main__":
    main()
