"""
Executive Summary Markdown to DOCX Converter
Converts EXECUTIVE_SUMMARY.md to a well-formatted Word document with proper tables
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_document_styles(doc):
    """Setup custom styles for the document with compact spacing"""
    styles = doc.styles
    
    # Heading 1
    try:
        h1 = styles['Heading 1']
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102)
        h1.paragraph_format.space_before = Pt(6)
        h1.paragraph_format.space_after = Pt(3)
    except KeyError:
        pass
    
    # Heading 2
    try:
        h2 = styles['Heading 2']
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 102, 204)
        h2.paragraph_format.space_before = Pt(4)
        h2.paragraph_format.space_after = Pt(2)
    except KeyError:
        pass
    
    # Heading 3
    try:
        h3 = styles['Heading 3']
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(51, 51, 51)
        h3.paragraph_format.space_before = Pt(3)
        h3.paragraph_format.space_after = Pt(2)
    except KeyError:
        pass


def clean_markdown_formatting(text):
    """Remove markdown formatting from text"""
    text = text.strip()
    # Remove bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Remove code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


def parse_table(lines, start_idx):
    """Parse markdown table and return table data and next line index"""
    i = start_idx
    table_lines = []
    
    # Collect all consecutive table lines
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|'):
            table_lines.append(line)
            i += 1
        else:
            break
    
    if len(table_lines) < 2:
        return None, i
    
    # Process table lines
    rows = []
    for line in table_lines:
        # Skip separator line (contains dashes)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        
        # Split by | and clean
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first and last elements
        cells = [c for c in cells if c]
        
        if cells:
            rows.append(cells)
    
    return rows, i


def add_table_to_doc(doc, rows):
    """Add a formatted table to the document with compact spacing"""
    if not rows or len(rows) < 1:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Add content with compact spacing
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= len(row.cells):
                continue
            
            cell = row.cells[j]
            cell.text = clean_markdown_formatting(cell_text)
            
            # Apply compact spacing to all cell paragraphs
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.0
                
                # Header row formatting
                if i == 0:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
    
    # Set reasonable column widths
    try:
        total_width = Inches(6.5)
        col_width = total_width / len(rows[0])
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
    except:
        pass


def process_inline_formatting(paragraph, text):
    """Process inline markdown formatting (bold, italic, code)"""
    # Pattern to match **bold**, *italic*, `code`
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        # Italic
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        # Code
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def convert_markdown_to_docx(md_file_path, docx_file_path):
    """Convert markdown file to DOCX with proper formatting"""
    print(f"Reading: {md_file_path}")
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    
    # Set compact margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    i = 0
    tables_processed = 0
    
    print("Processing content...")
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip horizontal rules
        if re.match(r'^[-*_]{3,}$', line.strip()):
            i += 1
            continue
        
        # Check for table
        if line.strip().startswith('|'):
            table_data, next_i = parse_table(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
                tables_processed += 1
                # Add minimal spacing after table
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                i = next_i
                continue
        
        # Headers
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue
        
        # Bullet points
        if re.match(r'^[•\-\*]\s+', line):
            text = re.sub(r'^[•\-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            i += 1
            continue
        
        # Checkboxes
        if re.match(r'^[☐✓✅❌]\s+', line):
            text = line.strip()
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            i += 1
            continue
        
        # Empty lines
        if not line.strip():
            i += 1
            continue
        
        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            process_inline_formatting(p, line)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
        
        i += 1
    
    print(f"Saving document to: {docx_file_path}")
    print(f"Total tables processed: {tables_processed}")
    doc.save(docx_file_path)
    print("✅ Conversion complete!")


def main():
    """Main function"""
    script_dir = Path(__file__).parent
    md_file = script_dir / "EXECUTIVE_SUMMARY.md"
    docx_file = script_dir / "EXECUTIVE_SUMMARY.docx"
    
    print("="*60)
    print("Executive Summary Markdown to DOCX Converter")
    print("="*60)
    
    if not md_file.exists():
        print(f"❌ Error: File not found: {md_file}")
        return
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
        print(f"\n✅ Successfully created: {docx_file}")
        print(f"   File size: {docx_file.stat().st_size / 1024:.2f} KB")
    except Exception as e:
        print(f"\n❌ Error during conversion: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
