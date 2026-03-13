"""
BRD Markdown to DOCX Converter
Converts BRDV3.7.md to a well-formatted Word document with tables, headers, and styling
Supports all BRD-specific formatting including user stories, business rules, and workflows
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    
    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._element.append(hyperlink)
    return hyperlink


def setup_document_styles(doc):
    """Setup custom styles for the document"""
    styles = doc.styles
    
    # Heading 1 style
    if 'Custom Heading 1' not in styles:
        h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.base_style = styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Calibri'
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 style
    if 'Custom Heading 2' not in styles:
        h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.base_style = styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Calibri'
        h2_font.size = Pt(16)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 102, 204)
    
    # Heading 3 style
    if 'Custom Heading 3' not in styles:
        h3_style = styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.base_style = styles['Heading 3']
        h3_font = h3_style.font
        h3_font.name = 'Calibri'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(51, 102, 153)
    
    # Code style with syntax highlighting
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        # Add light gray background (simplified approach)
        try:
            from docx.oxml.shared import qn
            from docx.oxml import parse_xml
            shading_elm = parse_xml(r'<w:shd {} w:fill="F8F8F8"/>'.format(qn('w:xmlns:w')))
            code_style._element.get_or_add_pPr().append(shading_elm)
        except:
            # If shading fails, continue without it
            pass
    
    # Proto code style
    if 'ProtoCode' not in styles:
        proto_style = styles.add_style('ProtoCode', WD_STYLE_TYPE.PARAGRAPH)
        proto_font = proto_style.font
        proto_font.name = 'Consolas'
        proto_font.size = Pt(8)
        proto_style.paragraph_format.left_indent = Inches(0.5)
        proto_style.paragraph_format.space_before = Pt(6)
        proto_style.paragraph_format.space_after = Pt(6)


def clean_markdown_formatting(text):
    """Remove markdown formatting from text"""
    if not text:
        return text
    
    # Remove bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Remove code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # Remove <br> tags
    text = text.replace('<br>', '\n')
    
    return text


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table into rows and columns - completely rewritten for robustness"""
    i = start_idx
    table_lines = []
    
    # Collect consecutive lines that look like table rows
    while i < len(lines):
        line = lines[i].strip()
        
        # Stop at empty line, heading, or horizontal rule
        if not line or line.startswith('#') or line.startswith('---') or line.startswith('==='):
            break
            
        # Check if line is a table row (has at least 2 pipes)
        if line.count('|') >= 2:
            table_lines.append(line)
            i += 1
        else:
            # Not a table line, stop
            break
    
    # Need at least 2 lines (header + data row minimum)
    if len(table_lines) < 2:
        return None, i
    
    # Process table lines into rows
    processed_rows = []
    
    for line in table_lines:
        # Check if separator line (all dashes, colons, pipes, spaces)
        if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
            continue  # Skip separator
        
        # Split by pipe
        parts = line.split('|')
        
        # Remove first/last empty parts (from leading/trailing pipes)
        if len(parts) > 0 and parts[0].strip() == '':
            parts = parts[1:]
        if len(parts) > 0 and parts[-1].strip() == '':
            parts = parts[:-1]
        
        # Clean each cell
        cells = [clean_markdown_formatting(p.strip()) for p in parts]
        
        # Only add non-empty rows
        if any(cell for cell in cells):
            processed_rows.append(cells)
    
    # Return None if no valid rows
    if not processed_rows:
        return None, i
    
    return processed_rows, i


def add_table_to_doc(doc, rows):
    """Add a formatted table to the document; drop 'Phase' column if present; set widths"""
    if not rows or len(rows) < 2:
        return
    
    # Detect 'Phase' column index and remove it from rows to save space
    headers = rows[0]
    phase_idx = None
    for idx, h in enumerate(headers):
        if h.strip().lower() == 'phase':
            phase_idx = idx
            break
    
    if phase_idx is not None:
        new_rows = []
        for r in rows:
            if len(r) > phase_idx:
                new_row = r[:phase_idx] + r[phase_idx+1:]
            else:
                new_row = r
            new_rows.append(new_row)
        rows = new_rows
        headers = rows[0]
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Add content
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= len(row.cells):  # Safety check
                continue
            cell = row.cells[j]
            cell.text = clean_markdown_formatting(cell_text)
            
            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    # Set column widths with priority to 'Requirement Description'
    try:
        col_names = [h.strip().lower() for h in headers]
        total_width = Inches(6.5)
        widths = [None] * len(headers)
        
        # Find requirement description column
        if 'requirement description' in col_names:
            rd_idx = col_names.index('requirement description')
        else:
            rd_idx = None
        
        # Default width allocation
        for idx in range(len(headers)):
            header_lower = headers[idx].strip().lower()
            if header_lower in ['fr id', 'id', 'nfr id', 'idra id', 'bfiu id', 'privacy id']:
                widths[idx] = Inches(0.9)
            elif header_lower == 'priority':
                widths[idx] = Inches(1.0)
            elif header_lower in ['acceptance criteria', 'notes']:
                widths[idx] = Inches(2.1)
            elif rd_idx is not None and idx == rd_idx:
                widths[idx] = Inches(3.0)  # Give more space to requirement description
            else:
                widths[idx] = Inches(1.5)
        
        # Normalize to total width
        current_total = sum(w.inches for w in widths)
        scale = 6.5 / current_total if current_total else 1.0
        widths = [Inches(w.inches * scale) for w in widths]
        
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width
    except Exception as e:
        # Fallback uniform widths
        print(f"Warning: Width calculation failed ({e}), using uniform widths")
        widths = [Inches(6.5 / len(headers))] * len(headers)
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width


import hashlib

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """Main conversion function"""
    
    print(f"Reading markdown file: {md_file_path}")
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    print("Creating Word document...")
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    i = 0
    in_code_block = False
    code_lines = []
    tables_processed = 0
    
    print("Processing content...")
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Page break hints: '[[[PAGEBREAK]]]' starts a new page
        if line.strip() == '[[[PAGEBREAK]]]':
            doc.add_page_break()
            i += 1
            continue

        # Image handling - detect ![alt](path) patterns
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            # Try to add image if file exists
            script_dir = Path(__file__).parent
            full_img_path = script_dir / img_path
            
            if full_img_path.exists():
                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(str(full_img_path), width=Inches(6.0))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add caption
                    if alt_text:
                        caption = doc.add_paragraph(f"*{alt_text}*")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in caption.runs:
                            run.font.italic = True
                            run.font.size = Pt(10)
                    
                    print(f"  Image added: {img_path}")
                except Exception as e:
                    print(f"  Warning: Could not add image {img_path}: {e}")
                    doc.add_paragraph(f"[Image: {alt_text}] - {img_path}")
            else:
                print(f"  Warning: Image not found: {full_img_path}")
                doc.add_paragraph(f"[Image placeholder: {alt_text}] - {img_path}")
            
            i += 1
            continue

        # Code blocks - enhanced with language detection
        code_match = re.match(r'^```(\w+)?', line)
        if code_match or line.startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    language = getattr(in_code_block, 'language', '') if hasattr(in_code_block, 'language') else ''
                    style = 'ProtoCode' if language == 'protobuf' else 'Code'
                    
                    # Add code with syntax awareness
                    code_content = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_content, style=style)
                    
                code_lines = []
                in_code_block = False
            else:
                # Start code block - capture language
                language = code_match.group(1) if code_match else ''
                in_code_block = type('obj', (object,), {'language': language})()
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(clean_markdown_formatting(line[2:]), level=1)
            i += 1
            continue
        
        if line.startswith('## '):
            doc.add_heading(clean_markdown_formatting(line[3:]), level=2)
            i += 1
            continue
        
        if line.startswith('### '):
            doc.add_heading(clean_markdown_formatting(line[4:]), level=3)
            i += 1
            continue
        
        if line.startswith('#### '):
            doc.add_heading(clean_markdown_formatting(line[5:]), level=4)
            i += 1
            continue
        
        # Tables: accept typical markdown tables and ignore accidental nested pipes in paragraphs
        if line.strip().startswith('|') and line.count('|') >= 2 and not line.strip().startswith('```'):
            rows, new_i = parse_markdown_table(lines, i)
            if rows and len(rows) >= 2 and len(rows[0]) >= 2:
                add_table_to_doc(doc, rows)
                doc.add_paragraph()  # Add spacing after table
                tables_processed += 1
                print(f"  Table {tables_processed} processed with {len(rows)} rows")
                i = new_i
                continue
            else:
                # Not a valid table, treat as regular text
                i += 1
                continue
        
        # Horizontal rules
        if line.startswith('---') or line.startswith('___'):
            p = doc.add_paragraph()
            p.paragraph_format.border_bottom = True
            i += 1
            continue
        
        # Bullet lists
        if line.startswith('- ') or line.startswith('* '):
            text = clean_markdown_formatting(line[2:])
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Numbered lists
        match = re.match(r'^\d+\.\s+(.+)', line)
        if match:
            text = clean_markdown_formatting(match.group(1))
            # Use bullet list to avoid Word auto-continuation (25., 26., etc.)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Bold paragraphs (likely subsection headers)
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(clean_markdown_formatting(line))
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue
        
        # Empty lines
        if not line.strip():
            i += 1
            continue
        
        # Regular paragraphs
        text = clean_markdown_formatting(line)
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
        
        i += 1
    
    print(f"Saving document to: {docx_file_path}")
    print(f"Total tables processed: {tables_processed}")
    doc.save(docx_file_path)
    print("✅ Conversion complete!")


def main():
    """Main function"""
    # File paths
    script_dir = Path(__file__).parent
    md_file = script_dir / "BRDV3.7.md"
    docx_file = script_dir / "BRDV3.7.docx"
    
    print("="*60)
    print("BRD Markdown to DOCX Converter")
    print("="*60)
    
    # Check if input file exists
    if not md_file.exists():
        print(f"❌ Error: File not found: {md_file}")
        return
    
    # Convert
    try:
        convert_markdown_to_docx(md_file, docx_file)
        print(f"\n✅ Successfully created: {docx_file}")
        print(f"   File size: {docx_file.stat().st_size / 1024:.2f} KB")
    except Exception as e:
        print(f"\n❌ Error during conversion: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
