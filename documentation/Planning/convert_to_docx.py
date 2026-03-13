#!/usr/bin/env python3
"""
Convert DetailedProjectPlan.md to DetailedProjectPlan.docx
Converts markdown to formatted Word document with proper styling.
"""

import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx library not found!")
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def setup_styles(doc):
    """Setup custom styles for the document."""
    styles = doc.styles
    
    # Heading styles are built-in, just modify them
    h1 = styles['Heading 1']
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)
    
    h2 = styles['Heading 2']
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 102, 204)
    
    h3 = styles['Heading 3']
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(51, 51, 51)
    
    h4 = styles['Heading 4']
    h4.font.size = Pt(12)
    h4.font.bold = True
    
    # Normal text style
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'

def parse_table(lines, start_idx):
    """Parse markdown table and return table data and end index."""
    table_lines = []
    idx = start_idx
    
    while idx < len(lines) and '|' in lines[idx]:
        table_lines.append(lines[idx])
        idx += 1
    
    if len(table_lines) < 2:
        return None, start_idx
    
    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Skip separator line
    rows = []
    for line in table_lines[2:]:
        if line.strip():
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
    
    return {'headers': headers, 'rows': rows}, idx

def add_table_to_doc(doc, table_data):
    """Add a table to the document."""
    if not table_data or not table_data['rows']:
        return
    
    headers = table_data['headers']
    rows = table_data['rows']
    
    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        # Bold header text
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = cell_data

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document."""
    print(f"Reading markdown file: {md_file}")
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    print("Creating Word document...")
    doc = Document()
    
    # Setup styles
    setup_styles(doc)
    
    # Add title page
    title = doc.add_heading('InsureTech Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Detailed Project Plan', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(f'\nGenerated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n').italic = True
    meta_para.add_run(f'Project Start: December 20, 2025\n').italic = True
    meta_para.add_run(f'Planning Period: Dec 20, 2025 - Aug 1, 2026\n').italic = True
    
    doc.add_page_break()
    
    # Process content
    idx = 0
    in_code_block = False
    code_lines = []
    
    while idx < len(lines):
        line = lines[idx]
        
        # Skip title if it's at the beginning
        if idx < 5 and line.startswith('# InsureTech'):
            idx += 1
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    code_para = doc.add_paragraph('\n'.join(code_lines))
                    code_para.style = 'Normal'
                    code_para_format = code_para.paragraph_format
                    code_para_format.left_indent = Inches(0.5)
                    for run in code_para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            idx += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        
        # Handle tables
        elif '|' in line and idx + 1 < len(lines) and '|' in lines[idx + 1]:
            table_data, new_idx = parse_table(lines, idx)
            if table_data:
                add_table_to_doc(doc, table_data)
                idx = new_idx
                continue
        
        # Handle bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line.strip()):
            text = re.sub(r'^\d+\. ', '', line.strip())
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Number')
        
        # Handle regular paragraphs
        elif line.strip():
            text = line.strip()
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            
            if text and not text.startswith('#'):
                para = doc.add_paragraph(text)
        
        idx += 1
    
    # Save document
    print(f"Saving Word document: {docx_file}")
    doc.save(docx_file)
    print("Conversion complete!")
    
    return docx_file

def main():
    """Main function."""
    script_dir = Path(__file__).parent
    md_file = script_dir / "DetailedProjectPlan.md"
    docx_file = script_dir / "DetailedProjectPlan.docx"
    
    print("=" * 70)
    print("Markdown to Word Converter")
    print("=" * 70)
    print()
    
    if not md_file.exists():
        print(f"Error: Markdown file not found: {md_file}")
        return 1
    
    try:
        output_file = convert_markdown_to_docx(md_file, docx_file)
        
        # Get file size
        size_kb = os.path.getsize(output_file) / 1024
        print(f"\nSuccess!")
        print(f"Output: {output_file}")
        print(f"Size: {size_kb:.2f} KB")
        
    except Exception as e:
        print(f"Error during conversion: {e}")
        import traceback
        traceback.print_exc()
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())
