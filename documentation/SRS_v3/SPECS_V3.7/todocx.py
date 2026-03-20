"""
SRS Markdown to DOCX Converter (SPECS_V3.7 version)
- Searches for any Markdown files starting with 'SRS' in the parent folder (SRS_v3)
- Generates a DOCX with the same name in the same folder as the Markdown
- Uses bullet lists for all list items to avoid Word auto-numbering continuation
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0000FF'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink


def setup_document_styles(doc):
    styles = doc.styles
    if 'Custom Heading 1' not in styles:
        h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.base_style = styles['Heading 1']
        h1_style.font.name = 'Calibri'; h1_style.font.size = Pt(18); h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 51, 102)
    if 'Custom Heading 2' not in styles:
        h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.base_style = styles['Heading 2']
        h2_style.font.name = 'Calibri'; h2_style.font.size = Pt(16); h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0, 102, 204)
    if 'Custom Heading 3' not in styles:
        h3_style = styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.base_style = styles['Heading 3']
        h3_style.font.name = 'Calibri'; h3_style.font.size = Pt(14); h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(51, 102, 153)
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'; code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    if 'ProtoCode' not in styles:
        proto_style = styles.add_style('ProtoCode', WD_STYLE_TYPE.PARAGRAPH)
        proto_style.font.name = 'Consolas'; proto_style.font.size = Pt(8)
        proto_style.paragraph_format.left_indent = Inches(0.5)
        proto_style.paragraph_format.space_before = Pt(6)
        proto_style.paragraph_format.space_after = Pt(6)


def clean_markdown_formatting(text):
    if not text:
        return text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = text.replace('<br>', '\n')
    return text


def parse_markdown_table(lines, start_idx):
    i = start_idx; table_lines = []
    while i < len(lines):
        line = lines[i].strip()
        if not line or line.startswith('#') or line.startswith('---') or line.startswith('==='):
            break
        if line.count('|') >= 2:
            table_lines.append(line); i += 1
        else:
            break
    if len(table_lines) < 2:
        return None, i
    processed_rows = []
    for line in table_lines:
        if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
            continue
        parts = line.split('|')
        if len(parts) > 0 and parts[0].strip() == '':
            parts = parts[1:]
        if len(parts) > 0 and parts[-1].strip() == '':
            parts = parts[:-1]
        cells = [clean_markdown_formatting(p.strip()) for p in parts]
        if any(cell for cell in cells):
            processed_rows.append(cells)
    if not processed_rows:
        return None, i
    return processed_rows, i


def add_table_to_doc(doc, rows):
    if not rows or len(rows) < 2:
        return
    headers = rows[0]
    phase_idx = None
    for idx, h in enumerate(headers):
        if h.strip().lower() == 'phase':
            phase_idx = idx; break
    if phase_idx is not None:
        new_rows = []
        for r in rows:
            new_row = r[:phase_idx] + r[phase_idx+1:] if len(r) > phase_idx else r
            new_rows.append(new_row)
        rows = new_rows; headers = rows[0]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= len(row.cells):
                continue
            cell = row.cells[j]
            cell.text = clean_markdown_formatting(cell_text)
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True; run.font.size = Pt(10)
                row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


def convert_markdown_to_docx(md_file_path, docx_file_path):
    print(f"Reading markdown file: {md_file_path}")
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    print("Creating Word document...")
    doc = Document()
    setup_document_styles(doc)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    i = 0; in_code_block = False; code_lines = []; tables_processed = 0
    print("Processing content...")
    while i < len(lines):
        line = lines[i].rstrip()
        if line.strip() == '[[[PAGEBREAK]]]':
            doc.add_page_break(); i += 1; continue
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt_text, img_path = img_match.group(1), img_match.group(2)
            script_dir = Path(__file__).parent
            # Attempt relative to this script and to the v3 root
            full_img_path = script_dir / img_path
            if not full_img_path.exists():
                full_img_path = (script_dir.parent / img_path)
            if full_img_path.exists():
                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(str(full_img_path), width=Inches(6.0))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if alt_text:
                        caption = doc.add_paragraph(f"*{alt_text}*")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in caption.runs:
                            run.font.italic = True; run.font.size = Pt(10)
                    print(f"  Image added: {img_path}")
                except Exception as e:
                    print(f"  Warning: Could not add image {img_path}: {e}")
                    doc.add_paragraph(f"[Image: {alt_text}] - {img_path}")
            else:
                print(f"  Warning: Image not found: {full_img_path}")
                doc.add_paragraph(f"[Image placeholder: {alt_text}] - {img_path}")
            i += 1; continue
        code_match = re.match(r'^```(\w+)?', line)
        if code_match or line.startswith('```'):
            if in_code_block:
                if code_lines:
                    language = getattr(in_code_block, 'language', '') if hasattr(in_code_block, 'language') else ''
                    style = 'ProtoCode' if language == 'protobuf' else 'Code'
                    code_content = '\n'.join(code_lines)
                    doc.add_paragraph(code_content, style=style)
                code_lines = []; in_code_block = False
            else:
                language = code_match.group(1) if code_match else ''
                in_code_block = type('obj', (object,), {'language': language})()
            i += 1; continue
        if in_code_block:
            code_lines.append(line); i += 1; continue
        if line.startswith('# '):
            doc.add_heading(clean_markdown_formatting(line[2:]), level=1); i += 1; continue
        if line.startswith('## '):
            doc.add_heading(clean_markdown_formatting(line[3:]), level=2); i += 1; continue
        if line.startswith('### '):
            doc.add_heading(clean_markdown_formatting(line[4:]), level=3); i += 1; continue
        if line.startswith('#### '):
            doc.add_heading(clean_markdown_formatting(line[5:]), level=4); i += 1; continue
        if line.strip().startswith('|') and line.count('|') >= 2 and not line.strip().startswith('```'):
            rows, new_i = parse_markdown_table(lines, i)
            if rows and len(rows) >= 2 and len(rows[0]) >= 2:
                add_table_to_doc(doc, rows); doc.add_paragraph(); tables_processed += 1
                print(f"  Table {tables_processed} processed with {len(rows)} rows")
                i = new_i; continue
            else:
                i += 1; continue
        if line.startswith('---') or line.startswith('___'):
            p = doc.add_paragraph(); p.paragraph_format.border_bottom = True; i += 1; continue
        if line.startswith('- ') or line.startswith('* '):
            text = clean_markdown_formatting(line[2:])
            doc.add_paragraph(text, style='List Bullet'); i += 1; continue
        match = re.match(r'^\d+\.\s+(.+)', line)
        if match:
            text = clean_markdown_formatting(match.group(1))
            doc.add_paragraph(text, style='List Bullet'); i += 1; continue
        if not line.strip():
            i += 1; continue
        text = clean_markdown_formatting(line)
        if text:
            p = doc.add_paragraph(text); p.paragraph_format.space_after = Pt(6)
        i += 1
    print(f"Saving document to: {docx_file_path}")
    print(f"Total tables processed: {tables_processed}")
    doc.save(docx_file_path)
    print("✅ Conversion complete!")


def main():
    script_dir = Path(__file__).parent
    v3_root = script_dir.parent
    print("="*60)
    print("SRS Markdown to DOCX Converter (SPECS_V3.7)")
    print("="*60)
    md_files = sorted(v3_root.glob("SRS*.md"))
    if not md_files:
        print(f"No SRS*.md files found in {v3_root}")
        return
    for md_file in md_files:
        docx_file = md_file.with_suffix('.docx')
        try:
            convert_markdown_to_docx(md_file, docx_file)
            print(f"\n✅ Successfully created: {docx_file}")
            print(f"   File size: {docx_file.stat().st_size / 1024:.2f} KB")
        except Exception as e:
            print(f"\n❌ Error during conversion for {md_file.name}: {e}")
            import traceback; traceback.print_exc()


if __name__ == "__main__":
    main()
