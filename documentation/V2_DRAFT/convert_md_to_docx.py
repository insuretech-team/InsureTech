#!/usr/bin/env python3
"""
Markdown to DOCX Converter for LabAid InsureTech SRS
Converts the V2_COMPLETE.md file to a professionally formatted DOCX document
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_custom_styles(doc):
    """Create custom styles for the document"""
    
    # Title style
    title_style = doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(20)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Subtitle style
    subtitle_style = doc.styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(14)
    subtitle_font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(8)
    
    # Header 1 style
    h1_style = doc.styles.add_style('Custom H1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Header 2 style
    h2_style = doc.styles.add_style('Custom H2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0x2f, 0x5f, 0x8f)
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)
    
    # Header 3 style
    h3_style = doc.styles.add_style('Custom H3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0x4f, 0x6f, 0x9f)
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(3)
    
    # Code style
    code_style = doc.styles.add_style('Custom Code', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    
    # Table header style
    table_header_style = doc.styles.add_style('Table Header', WD_STYLE_TYPE.PARAGRAPH)
    table_header_font = table_header_style.font
    table_header_font.name = 'Calibri'
    table_header_font.size = Pt(11)
    table_header_font.bold = True
    table_header_font.color.rgb = RGBColor(0xff, 0xff, 0xff)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def parse_table(table_text):
    """Parse markdown table to list of rows"""
    lines = table_text.strip().split('\n')
    # Remove empty lines
    lines = [line.strip() for line in lines if line.strip()]
    
    # Find table rows (lines with |)
    table_rows = []
    for line in lines:
        if '|' in line and not line.startswith('|-'):
            # Split by | and clean up
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty first/last cells if they exist
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            if cells:
                table_rows.append(cells)
    
    return table_rows

def create_table(doc, table_data):
    """Create a formatted table"""
    if not table_data or len(table_data) < 2:
        return
    
    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Format header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(table_data[0]):
        if i < len(header_cells):
            header_cells[i].text = header_text
            # Apply header formatting
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set cell background color
            cell_xml_element = header_cells[i]._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), "1f4e79")
            table_cell_properties.append(shade_obj)
    
    # Format data rows
    for row_idx in range(1, len(table_data)):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_text in enumerate(table_data[row_idx]):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = cell_text
                # Center align ID columns
                if col_idx == 0:  # First column (usually ID)
                    for paragraph in row_cells[col_idx].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adjust column widths
    if len(table_data[0]) >= 3:
        # Typical requirement table format
        table.columns[0].width = Inches(0.8)  # ID column
        table.columns[1].width = Inches(4.5)  # Description column
        if len(table.columns) > 2:
            table.columns[2].width = Inches(0.8)  # Priority column

def process_markdown_content(content):
    """Process markdown content and convert to DOCX"""
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Create custom styles
    create_custom_styles(doc)
    
    # Split content into lines
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle main title
        if line.startswith('# ') and 'System Requirements Specification' in line:
            title_para = doc.add_paragraph(line[2:], style='Custom Title')
            i += 1
            continue
        
        # Handle project info
        if line.startswith('**Project:**'):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            i += 1
            continue
        
        # Handle version and date info
        if line.startswith('**Version:**') or line.startswith('**Date:**') or line.startswith('**Document Classification:**'):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            i += 1
            continue
        
        # Handle headers
        if line.startswith('## '):
            # Check for page break before major sections
            if any(section in line for section in ['External Interface', 'Non-Functional', 'Data Model', 'Security', 'Performance', 'AML/CFT', 'Operational', 'Acceptance', 'Traceability']):
                add_page_break(doc)
            doc.add_paragraph(line[3:], style='Custom H1')
            i += 1
            continue
        
        if line.startswith('### '):
            doc.add_paragraph(line[4:], style='Custom H2')
            i += 1
            continue
        
        if line.startswith('#### '):
            doc.add_paragraph(line[5:], style='Custom H3')
            i += 1
            continue
        
        # Handle tables
        if '|' in line and not line.startswith('#'):
            # Collect table lines
            table_lines = []
            while i < len(lines) and ('|' in lines[i] or lines[i].strip() == ''):
                if lines[i].strip():
                    table_lines.append(lines[i])
                i += 1
            
            # Process table
            if table_lines:
                table_data = parse_table('\n'.join(table_lines))
                if table_data:
                    create_table(doc, table_data)
                    doc.add_paragraph()  # Add spacing after table
            continue
        
        # Handle code blocks
        if line.startswith('```'):
            i += 1  # Skip opening ```
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip closing ```
            
            # Add code block
            if code_lines:
                code_text = '\n'.join(code_lines)
                doc.add_paragraph(code_text, style='Custom Code')
                doc.add_paragraph()  # Add spacing after code
            continue
        
        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.', line):
            para = doc.add_paragraph(line, style='List Number')
            i += 1
            continue
        
        # Handle horizontal rules
        if line.startswith('---'):
            para = doc.add_paragraph()
            para.add_run('_' * 80)
            i += 1
            continue
        
        # Handle bold text in paragraphs
        if '**' in line:
            para = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for j, part in enumerate(parts):
                if j % 2 == 0:
                    para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.bold = True
            i += 1
            continue
        
        # Regular paragraph
        if line:
            doc.add_paragraph(line)
        
        i += 1
    
    return doc

def main():
    """Main conversion function"""
    input_file = r"C:\_DEV\GO\InsureTech\V2.2_COMPLETE.md"
    output_file = r"C:\_DEV\GO\InsureTech\LabAid_InsureTech_SRS_V2.2_COMPLETE.docx"
    
    try:
        # Read markdown content
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Process content and create document
        doc = process_markdown_content(content)
        
        # Add header
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "LabAid InsureTech Platform - System Requirements Specification v2.0"
        header_para.style = doc.styles['Header']
        
        # Add footer with page numbers
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Page "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        doc.save(output_file)
        print(f"✅ Successfully converted {input_file} to {output_file}")
        print(f"📄 Output file size: {os.path.getsize(output_file)} bytes")
        
    except FileNotFoundError:
        print(f"❌ Error: Could not find input file {input_file}")
    except Exception as e:
        print(f"❌ Error during conversion: {str(e)}")

if __name__ == "__main__":
    # Install required package if not available
    try:
        import docx
    except ImportError:
        print("Installing python-docx package...")
        os.system("pip install python-docx")
        import docx
    
    main()